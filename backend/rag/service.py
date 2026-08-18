"""Shared Gemini and FAISS document retrieval services for LegalSimple."""

import io
import logging
import re
import threading
from collections import Counter
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from uuid import uuid4

import pymupdf
from docx import Document as DocxDocument
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend import config
from backend.rag.prompts import (
    DOCUMENT_ANSWER_SYSTEM_PROMPT,
    DOCUMENT_COMPARISON_SYSTEM_PROMPT,
    DOCUMENT_SIMPLIFIER_RETRIEVAL_QUERY,
    DOCUMENT_SIMPLIFIER_SYSTEM_PROMPT,
    GENERAL_CHAT_SYSTEM_PROMPT,
    QUESTION_REWRITE_SYSTEM_PROMPT,
    RISK_ANALYZER_RETRIEVAL_QUERY,
    RISK_ANALYZER_SYSTEM_PROMPT,
)
from backend.rag.schemas import (
    ChatHistoryMessage,
    ComparisonSource,
    DocumentComparisonItem,
    DocumentComparisonSummary,
    DocumentSource,
    DocumentUploadResponse,
    GeneratedDocumentComparison,
    GeneratedRiskAnalysis,
)


logger = logging.getLogger(__name__)
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
RISK_ANALYSIS_DISCLAIMER = (
    "This risk analysis is informational and is not legal advice."
)
COMPARISON_DISCLAIMER = (
    "This comparison is informational and is not legal advice. "
    "The original documents remain authoritative."
)
IMPORTANT_VALUE_PATTERN = re.compile(
    r"(?:\b(?:BDT|USD|EUR|GBP|Tk)\.?\s*[\d,]+(?:\.\d+)?\b|"
    r"[$€£৳]\s*[\d,]+(?:\.\d+)?|"
    r"\b\d+(?:\.\d+)?\s*%|"
    r"\b\d{1,4}[/-]\d{1,2}[/-]\d{1,4}\b|"
    r"\b(?:January|February|March|April|May|June|July|August|September|"
    r"October|November|December)\s+\d{1,2}(?:st|nd|rd|th)?(?:,\s*\d{4})?\b|"
    r"\b\d+(?:\.\d+)?\s+(?:business\s+)?(?:days?|weeks?|months?|years?|hours?)\b|"
    r"\b\d[\d,]*(?:\.\d+)?\b)",
    flags=re.IGNORECASE,
)
IMPORTANT_LEGAL_TERMS = (
    "amount",
    "payment",
    "fee",
    "penalty",
    "interest",
    "deadline",
    "notice",
    "terminate",
    "termination",
    "renewal",
    "default",
    "liability",
    "indemn",
    "obligation",
    "right",
)


class RAGConfigurationError(RuntimeError):
    """Raised when required Gemini configuration is unavailable."""


class DocumentProcessingError(ValueError):
    """Raised when an uploaded document cannot be parsed or indexed."""


class UnknownDocumentError(LookupError):
    """Raised when a document identifier has no in-memory FAISS store."""


class GeminiRequestError(RuntimeError):
    """Raised when Gemini cannot complete an embedding or generation request."""


@dataclass(frozen=True)
class DocumentStore:
    """The in-memory FAISS index and metadata for one uploaded document."""

    filename: str
    chunk_count: int
    chunks: tuple[Document, ...]
    vector_store: FAISS


@dataclass(frozen=True)
class ComparisonCandidate:
    """One deterministic cross-document match prepared for Gemini."""

    comparison_id: str
    change_type: str
    original: Document | None
    revised: Document | None
    text_similarity: float
    semantic_distance: float | None
    original_only_values: tuple[str, ...]
    revised_only_values: tuple[str, ...]
    priority: float


class RAGService:
    """Parse, index, retrieve, and generate answers from uploaded documents."""

    def __init__(self) -> None:
        self._llm: ChatGoogleGenerativeAI | None = None
        self._embeddings: GoogleGenerativeAIEmbeddings | None = None
        self._document_stores: dict[str, DocumentStore] = {}
        self._stores_lock = threading.RLock()
        self._models_lock = threading.Lock()
        self._text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            length_function=len,
        )

    def process_document(
        self,
        filename: str,
        file_bytes: bytes,
        document_label: str | None = None,
    ) -> DocumentUploadResponse:
        """Parse and index one supported document in an isolated FAISS store."""
        safe_filename = Path(filename).name.strip()
        extension = Path(safe_filename).suffix.lower()

        if not safe_filename or extension not in SUPPORTED_EXTENSIONS:
            raise DocumentProcessingError("Choose a PDF, DOCX, or TXT document.")
        if not file_bytes:
            raise DocumentProcessingError("The selected document is empty.")
        if len(file_bytes) > config.MAX_UPLOAD_BYTES:
            raise DocumentProcessingError("The selected document is larger than 10 MB.")

        try:
            documents = self._parse_document(safe_filename, extension, file_bytes)
            chunks = self._text_splitter.split_documents(documents)
        except DocumentProcessingError:
            raise
        except Exception as error:
            logger.exception("Document parsing failed for %s", safe_filename)
            raise DocumentProcessingError("The document could not be read.") from error

        chunks = [chunk for chunk in chunks if chunk.page_content.strip()]
        if not chunks:
            raise DocumentProcessingError("No extractable text was found in the document.")

        for chunk_index, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = chunk_index
            if document_label:
                chunk.metadata["document_label"] = document_label

        _, embeddings = self._ensure_models()
        try:
            vector_store = FAISS.from_documents(documents=chunks, embedding=embeddings)
        except Exception as error:
            logger.exception(
                "Gemini embedding or FAISS indexing failed for %s", safe_filename
            )
            raise GeminiRequestError(
                "The document could not be processed by the AI service."
            ) from error

        document_id = uuid4().hex
        with self._stores_lock:
            self._document_stores[document_id] = DocumentStore(
                filename=safe_filename,
                chunk_count=len(chunks),
                chunks=tuple(chunks),
                vector_store=vector_store,
            )

        logger.info(
            "Indexed %s as %s with %d chunks",
            safe_filename,
            document_id,
            len(chunks),
        )
        return DocumentUploadResponse(
            status="ready",
            document_id=document_id,
            filename=safe_filename,
            chunks=len(chunks),
        )

    def remove_document(self, document_id: str) -> None:
        """Remove one document's in-memory FAISS store."""
        with self._stores_lock:
            removed_store = self._document_stores.pop(document_id, None)
        if removed_store is None:
            raise UnknownDocumentError("The attached document is no longer available.")

    def chat(
        self,
        message: str,
        document_id: str | None,
        history: list[ChatHistoryMessage],
    ) -> tuple[str, list[DocumentSource]]:
        """Generate either a general answer or a document-grounded RAG answer."""
        llm, _ = self._ensure_models()
        langchain_history = self._to_langchain_history(history)

        try:
            if document_id:
                return self._chat_with_document(
                    message, document_id, langchain_history, llm
                )
            return self._general_chat(message, langchain_history, llm), []
        except UnknownDocumentError:
            raise
        except Exception as error:
            logger.exception("Gemini generation failed")
            raise GeminiRequestError(
                "The Legal AI Companion could not generate a response."
            ) from error

    def simplify_document(
        self, document_id: str
    ) -> tuple[str, list[DocumentSource]]:
        """Generate a plain-language guide from broadly retrieved document chunks."""
        llm, _ = self._ensure_models()

        try:
            document_store = self._get_document_store(document_id)
            retrieved_documents = self._retrieve_documents(
                document_id,
                DOCUMENT_SIMPLIFIER_RETRIEVAL_QUERY,
                config.SIMPLIFIER_RETRIEVAL_COUNT,
            )
            # Similarity search returns relevance order. Restoring source order makes
            # connected clauses and exceptions easier for Gemini to interpret.
            retrieved_documents.sort(
                key=lambda document: document.metadata.get("chunk_index", 0)
            )
            context = self._format_context(retrieved_documents)
            simplifier_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", DOCUMENT_SIMPLIFIER_SYSTEM_PROMPT),
                    ("human", "Create a plain-language explanation of this document."),
                ]
            )
            if len(retrieved_documents) == document_store.chunk_count:
                retrieval_coverage = "All document chunks were retrieved."
            else:
                retrieval_coverage = (
                    f"{len(retrieved_documents)} of {document_store.chunk_count} "
                    "document chunks were retrieved by relevance."
                )
            simplification = (simplifier_prompt | llm | StrOutputParser()).invoke(
                {
                    "context": context,
                    "retrieval_coverage": retrieval_coverage,
                }
            )
        except UnknownDocumentError:
            raise
        except Exception as error:
            logger.exception("Gemini document simplification failed")
            raise GeminiRequestError(
                "The document could not be simplified by the AI service."
            ) from error

        return self._prepare_answer_for_display(
            simplification
        ), self._collect_sources(retrieved_documents)

    def analyze_document_risk(
        self, document_id: str
    ) -> tuple[GeneratedRiskAnalysis, list[DocumentSource]]:
        """Classify apparent risk using relevant chunks from one indexed document."""
        llm, _ = self._ensure_models()

        try:
            document_store = self._get_document_store(document_id)
            retrieved_documents = self._retrieve_documents(
                document_id,
                RISK_ANALYZER_RETRIEVAL_QUERY,
                config.RISK_ANALYZER_RETRIEVAL_COUNT,
            )
            # Source order helps Gemini interpret clauses together with nearby
            # qualifications and exceptions.
            retrieved_documents.sort(
                key=lambda document: document.metadata.get("chunk_index", 0)
            )
            context = self._format_context(retrieved_documents)
            if len(retrieved_documents) == document_store.chunk_count:
                retrieval_coverage = "All document chunks were retrieved."
            else:
                retrieval_coverage = (
                    f"{len(retrieved_documents)} of {document_store.chunk_count} "
                    "document chunks were retrieved by relevance."
                )

            risk_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", RISK_ANALYZER_SYSTEM_PROMPT),
                    ("human", "Assess the overall apparent risk of this document."),
                ]
            )
            structured_llm = llm.with_structured_output(
                GeneratedRiskAnalysis,
                method="json_schema",
            )
            analysis = (risk_prompt | structured_llm).invoke(
                {
                    "context": context,
                    "retrieval_coverage": retrieval_coverage,
                }
            )
            if not isinstance(analysis, GeneratedRiskAnalysis):
                raise TypeError("Gemini returned an unexpected risk analysis format.")

            explanation = self._prepare_answer_for_display(analysis.explanation)
            if RISK_ANALYSIS_DISCLAIMER.lower() not in explanation.lower():
                explanation = f"{explanation}\n\n{RISK_ANALYSIS_DISCLAIMER}"
            analysis = analysis.model_copy(update={"explanation": explanation})
        except UnknownDocumentError:
            raise
        except Exception as error:
            logger.exception("Gemini document risk analysis failed")
            raise GeminiRequestError(
                "The document could not be analyzed by the AI service."
            ) from error

        return analysis, self._collect_sources(retrieved_documents)

    def compare_documents(
        self,
        original_document_id: str,
        revised_document_id: str,
    ) -> tuple[
        str,
        DocumentComparisonSummary,
        list[DocumentComparisonItem],
        str,
        str,
    ]:
        """Cross-match two indexed PDFs and explain their supported differences."""
        llm, _ = self._ensure_models()

        try:
            original_store = self._get_document_store(original_document_id)
            revised_store = self._get_document_store(revised_document_id)

            # Match with existing FAISS vectors before asking Gemini to explain results.
            all_candidates = self._create_comparison_candidates(
                original_store, revised_store
            )
            selected_candidates = self._select_comparison_candidates(all_candidates)
            coverage_note = self._comparison_coverage_note(
                len(selected_candidates), len(all_candidates)
            )
            comparison_context = self._format_comparison_context(selected_candidates)

            comparison_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", DOCUMENT_COMPARISON_SYSTEM_PROMPT),
                    (
                        "human",
                        "Explain the prepared document comparison candidates in the "
                        "required structured format.",
                    ),
                ]
            )
            structured_llm = llm.with_structured_output(
                GeneratedDocumentComparison,
                method="json_schema",
            )
            generated_comparison = (comparison_prompt | structured_llm).invoke(
                {
                    "candidate_count": len(selected_candidates),
                    "coverage_note": coverage_note,
                    "comparison_context": comparison_context,
                }
            )
            if not isinstance(generated_comparison, GeneratedDocumentComparison):
                raise TypeError("Gemini returned an unexpected comparison format.")

            generated_items = {
                item.comparison_id: item for item in generated_comparison.items
            }
            expected_ids = {
                candidate.comparison_id for candidate in selected_candidates
            }
            if (
                len(generated_items) != len(generated_comparison.items)
                or set(generated_items) != expected_ids
            ):
                raise TypeError(
                    "Gemini did not return every prepared comparison candidate exactly once."
                )

            comparison_items: list[DocumentComparisonItem] = []
            for candidate in selected_candidates:
                generated_item = generated_items[candidate.comparison_id]

                # Source text and page metadata come from parsed chunks, not Gemini.
                comparison_items.append(
                    DocumentComparisonItem(
                        comparison_id=candidate.comparison_id,
                        section_title=self._prepare_answer_for_display(
                            generated_item.section_title
                        ),
                        change_type=candidate.change_type,
                        original_text=(
                            candidate.original.page_content.strip()
                            if candidate.original is not None
                            else None
                        ),
                        revised_text=(
                            candidate.revised.page_content.strip()
                            if candidate.revised is not None
                            else None
                        ),
                        explanation=self._prepare_answer_for_display(
                            generated_item.explanation
                        ),
                        potential_significance=self._prepare_answer_for_display(
                            generated_item.potential_significance
                        ),
                        important=(
                            generated_item.important
                            if candidate.change_type != "unchanged"
                            else False
                        ),
                        original_source=self._comparison_source(
                            candidate.original, "A"
                        ),
                        revised_source=self._comparison_source(
                            candidate.revised, "B"
                        ),
                    )
                )

            summary = self._summarize_comparison_items(comparison_items)
            overall_summary = self._prepare_answer_for_display(
                generated_comparison.overall_summary
            )
        except UnknownDocumentError:
            raise
        except Exception as error:
            logger.exception("Gemini document comparison failed")
            raise GeminiRequestError(
                "The documents could not be compared by the AI service."
            ) from error

        return (
            overall_summary,
            summary,
            comparison_items,
            coverage_note,
            COMPARISON_DISCLAIMER,
        )

    def _create_comparison_candidates(
        self,
        original_store: DocumentStore,
        revised_store: DocumentStore,
    ) -> list[ComparisonCandidate]:
        """Match chunks in both directions, then retain unmatched additions/removals."""
        cross_edges: dict[tuple[int, int], dict[str, object]] = {}
        self._add_cross_document_edges(
            original_store,
            revised_store,
            cross_edges,
            direction="A_to_B",
        )
        self._add_cross_document_edges(
            revised_store,
            original_store,
            cross_edges,
            direction="B_to_A",
        )

        ranked_edges: list[tuple[bool, float, float, int, int]] = []
        for (original_index, revised_index), edge in cross_edges.items():
            original_text = original_store.chunks[original_index].page_content
            revised_text = revised_store.chunks[revised_index].page_content
            text_similarity = self._text_similarity(original_text, revised_text)
            semantic_distance = float(edge["distance"])
            directions = edge["directions"]
            is_mutual = isinstance(directions, set) and len(directions) == 2

            should_match = (
                text_similarity == 1.0
                or semantic_distance <= config.COMPARISON_MAX_DISTANCE
                or (
                    is_mutual
                    and semantic_distance <= config.COMPARISON_MUTUAL_MAX_DISTANCE
                )
                or text_similarity >= config.COMPARISON_MIN_TEXT_SIMILARITY
            )
            if should_match:
                # Exact or strongly overlapping text wins before semantic distance.
                match_rank = semantic_distance - (0.20 * text_similarity)
                ranked_edges.append(
                    (
                        text_similarity == 1.0,
                        match_rank,
                        text_similarity,
                        original_index,
                        revised_index,
                    )
                )

        ranked_edges.sort(
            key=lambda edge: (
                not edge[0],
                edge[1],
                -edge[2],
                edge[3],
                edge[4],
            )
        )
        matched_original: set[int] = set()
        matched_revised: set[int] = set()
        matches: list[tuple[int, int, float, float]] = []
        for _, _, text_similarity, original_index, revised_index in ranked_edges:
            if (
                original_index in matched_original
                or revised_index in matched_revised
            ):
                continue
            matched_original.add(original_index)
            matched_revised.add(revised_index)
            semantic_distance = float(
                cross_edges[(original_index, revised_index)]["distance"]
            )
            matches.append(
                (
                    original_index,
                    revised_index,
                    text_similarity,
                    semantic_distance,
                )
            )

        candidates: list[ComparisonCandidate] = []
        next_id = 1
        for original_index, revised_index, text_similarity, distance in sorted(matches):
            original = original_store.chunks[original_index]
            revised = revised_store.chunks[revised_index]
            change_type = (
                "unchanged"
                if self._normalize_comparison_text(original.page_content)
                == self._normalize_comparison_text(revised.page_content)
                else "modified"
            )
            candidates.append(
                self._make_comparison_candidate(
                    next_id,
                    change_type,
                    original,
                    revised,
                    text_similarity,
                    distance,
                )
            )
            next_id += 1

        for original_index, original in enumerate(original_store.chunks):
            if original_index in matched_original:
                continue
            candidates.append(
                self._make_comparison_candidate(
                    next_id, "removed", original, None, 0.0, None
                )
            )
            next_id += 1

        for revised_index, revised in enumerate(revised_store.chunks):
            if revised_index in matched_revised:
                continue
            candidates.append(
                self._make_comparison_candidate(
                    next_id, "added", None, revised, 0.0, None
                )
            )
            next_id += 1

        return candidates

    @staticmethod
    def _add_cross_document_edges(
        source_store: DocumentStore,
        target_store: DocumentStore,
        edges: dict[tuple[int, int], dict[str, object]],
        direction: str,
    ) -> None:
        """Search every source vector against the other document's FAISS index."""
        retrieval_count = min(
            config.COMPARISON_MATCH_COUNT, target_store.chunk_count
        )
        for source_index in range(source_store.chunk_count):
            source_vector = source_store.vector_store.index.reconstruct(source_index)
            matches = target_store.vector_store.similarity_search_with_score_by_vector(
                source_vector.tolist(),
                k=retrieval_count,
            )
            for matched_document, raw_distance in matches:
                target_index = matched_document.metadata.get("chunk_index")
                if not isinstance(target_index, int):
                    continue
                if direction == "A_to_B":
                    edge_key = (source_index, target_index)
                else:
                    edge_key = (target_index, source_index)

                edge = edges.setdefault(
                    edge_key,
                    {"distance": float(raw_distance), "directions": set()},
                )
                edge["distance"] = min(
                    float(edge["distance"]), float(raw_distance)
                )
                directions = edge["directions"]
                if isinstance(directions, set):
                    directions.add(direction)

    def _make_comparison_candidate(
        self,
        candidate_number: int,
        change_type: str,
        original: Document | None,
        revised: Document | None,
        text_similarity: float,
        semantic_distance: float | None,
    ) -> ComparisonCandidate:
        original_text = original.page_content if original is not None else ""
        revised_text = revised.page_content if revised is not None else ""
        original_values, revised_values = self._different_significant_values(
            original_text, revised_text
        )
        combined_text = f"{original_text}\n{revised_text}".casefold()
        important_term_count = sum(
            term in combined_text for term in IMPORTANT_LEGAL_TERMS
        )
        priority = {
            "added": 80.0,
            "removed": 80.0,
            "modified": 60.0,
            "unchanged": 0.0,
        }[change_type]
        if original_values or revised_values:
            priority += 30.0
        priority += min(important_term_count * 2.0, 16.0)
        if change_type == "modified":
            priority += (1.0 - text_similarity) * 10.0

        return ComparisonCandidate(
            comparison_id=f"comparison-{candidate_number}",
            change_type=change_type,
            original=original,
            revised=revised,
            text_similarity=text_similarity,
            semantic_distance=semantic_distance,
            original_only_values=original_values,
            revised_only_values=revised_values,
            priority=priority,
        )

    @staticmethod
    def _select_comparison_candidates(
        candidates: list[ComparisonCandidate],
    ) -> list[ComparisonCandidate]:
        changed_candidates = sorted(
            (
                candidate
                for candidate in candidates
                if candidate.change_type != "unchanged"
            ),
            key=lambda candidate: (-candidate.priority, candidate.comparison_id),
        )
        selected = changed_candidates[: config.COMPARISON_MAX_ITEMS]
        remaining_capacity = config.COMPARISON_MAX_ITEMS - len(selected)
        if remaining_capacity > 0:
            # A few unchanged matches help confirm that corresponding text was found.
            unchanged_candidates = [
                candidate
                for candidate in candidates
                if candidate.change_type == "unchanged"
            ]
            selected.extend(
                unchanged_candidates[
                    : min(config.COMPARISON_UNCHANGED_ITEMS, remaining_capacity)
                ]
            )
        return selected

    @staticmethod
    def _normalize_comparison_text(text: str) -> str:
        return re.sub(r"\s+", " ", text).strip().casefold()

    @classmethod
    def _text_similarity(cls, original_text: str, revised_text: str) -> float:
        original_normalized = cls._normalize_comparison_text(original_text)
        revised_normalized = cls._normalize_comparison_text(revised_text)
        return SequenceMatcher(
            None, original_normalized, revised_normalized, autojunk=False
        ).ratio()

    @staticmethod
    def _different_significant_values(
        original_text: str, revised_text: str
    ) -> tuple[tuple[str, ...], tuple[str, ...]]:
        # Literal checks keep small changes such as 30 days to 60 days visible.
        original_values = IMPORTANT_VALUE_PATTERN.findall(original_text)
        revised_values = IMPORTANT_VALUE_PATTERN.findall(revised_text)
        original_remaining = Counter(value.casefold() for value in original_values)
        revised_remaining = Counter(value.casefold() for value in revised_values)
        shared_values = original_remaining & revised_remaining
        original_remaining -= shared_values
        revised_remaining -= shared_values

        def collect_remaining_values(
            values: list[str], remaining_counts: Counter[str]
        ) -> tuple[str, ...]:
            remaining_values: list[str] = []
            for value in values:
                normalized_value = value.casefold()
                if remaining_counts[normalized_value] <= 0:
                    continue
                remaining_values.append(value)
                remaining_counts[normalized_value] -= 1
            return tuple(remaining_values)

        original_only = collect_remaining_values(
            original_values, original_remaining
        )
        revised_only = collect_remaining_values(revised_values, revised_remaining)
        return original_only, revised_only

    @staticmethod
    def _format_comparison_context(
        candidates: list[ComparisonCandidate],
    ) -> str:
        formatted_candidates: list[str] = []
        for candidate in candidates:
            original_label = RAGService._comparison_document_label(
                candidate.original, "Document A"
            )
            revised_label = RAGService._comparison_document_label(
                candidate.revised, "Document B"
            )
            original_text = (
                candidate.original.page_content.strip()
                if candidate.original is not None
                else "[No corresponding passage in Document A]"
            )
            revised_text = (
                candidate.revised.page_content.strip()
                if candidate.revised is not None
                else "[No corresponding passage in Document B]"
            )
            original_values = ", ".join(candidate.original_only_values) or "none"
            revised_values = ", ".join(candidate.revised_only_values) or "none"
            semantic_note = (
                f"{candidate.semantic_distance:.4f} (lower means closer)"
                if candidate.semantic_distance is not None
                else "not applicable because one side is unmatched"
            )
            formatted_candidates.append(
                f'<comparison_candidate id="{candidate.comparison_id}" '
                f'preliminary_change_type="{candidate.change_type}">\n'
                f"Original source: {original_label}\n"
                f"BEGIN ORIGINAL EXCERPT\n{original_text}\nEND ORIGINAL EXCERPT\n"
                f"Revised source: {revised_label}\n"
                f"BEGIN REVISED EXCERPT\n{revised_text}\nEND REVISED EXCERPT\n"
                "DETERMINISTIC TEXT CHECK\n"
                f"Normalized text similarity: {candidate.text_similarity:.4f}\n"
                f"Semantic distance: {semantic_note}\n"
                f"Literal values only in original: {original_values}\n"
                f"Literal values only in revised: {revised_values}\n"
                "END DETERMINISTIC TEXT CHECK\n"
                "</comparison_candidate>"
            )
        return "\n\n".join(formatted_candidates)

    @staticmethod
    def _comparison_document_label(
        document: Document | None, default_label: str
    ) -> str:
        if document is None:
            return f"{default_label}, no source passage"
        filename = document.metadata.get("filename", "Uploaded PDF")
        page_number = document.metadata.get("page_number")
        label = f"{default_label}: {filename}"
        if isinstance(page_number, int):
            label += f", page {page_number}"
        return label

    @staticmethod
    def _comparison_source(
        document: Document | None, document_label: str
    ) -> ComparisonSource | None:
        if document is None:
            return None
        filename = document.metadata.get("filename")
        if not isinstance(filename, str) or not filename:
            return None
        page_number = document.metadata.get("page_number")
        if not isinstance(page_number, int):
            page_number = None
        return ComparisonSource(
            document_label=document_label,
            filename=filename,
            page_number=page_number,
        )

    @staticmethod
    def _summarize_comparison_items(
        items: list[DocumentComparisonItem],
    ) -> DocumentComparisonSummary:
        counts = Counter(item.change_type for item in items)
        return DocumentComparisonSummary(
            modifications=counts["modified"],
            additions=counts["added"],
            removals=counts["removed"],
            unchanged=counts["unchanged"],
        )

    @staticmethod
    def _comparison_coverage_note(selected_count: int, total_count: int) -> str:
        if selected_count == total_count:
            return (
                f"All {total_count} chunk-level comparison candidates were included."
            )
        return (
            f"Showing {selected_count} of {total_count} chunk-level comparison "
            "candidates, prioritized for additions, removals, changed literal "
            "values, and legally significant wording."
        )

    def _ensure_models(
        self,
    ) -> tuple[ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings]:
        if not config.GOOGLE_API_KEY:
            raise RAGConfigurationError(
                "The Gemini API key is not configured on the backend."
            )

        with self._models_lock:
            try:
                if self._llm is None:
                    self._llm = ChatGoogleGenerativeAI(
                        model=config.GEMINI_GENERATION_MODEL,
                        temperature=1.0,
                        max_retries=2,
                    )
                if self._embeddings is None:
                    self._embeddings = GoogleGenerativeAIEmbeddings(
                        model=config.GEMINI_EMBEDDING_MODEL,
                    )
            except Exception as error:
                logger.exception("Gemini clients could not be initialized")
                raise GeminiRequestError(
                    "The Legal AI Companion could not initialize the AI service."
                ) from error

        assert self._llm is not None and self._embeddings is not None
        return self._llm, self._embeddings

    def _parse_document(
        self, filename: str, extension: str, file_bytes: bytes
    ) -> list[Document]:
        if extension == ".pdf":
            documents = self._parse_pdf(filename, file_bytes)
        elif extension == ".docx":
            documents = self._parse_docx(filename, file_bytes)
        else:
            documents = self._parse_txt(filename, file_bytes)

        extracted_characters = sum(len(document.page_content) for document in documents)
        if extracted_characters > config.MAX_EXTRACTED_CHARACTERS:
            raise DocumentProcessingError(
                "The document contains too much extracted text for this version of LegalSimple."
            )
        if not documents or extracted_characters == 0:
            raise DocumentProcessingError("No extractable text was found in the document.")
        return documents

    @staticmethod
    def _parse_pdf(filename: str, file_bytes: bytes) -> list[Document]:
        documents: list[Document] = []
        with pymupdf.open(stream=file_bytes, filetype="pdf") as pdf_document:
            for page_index, page in enumerate(pdf_document):
                page_text = page.get_text("text").strip()
                if page_text:
                    documents.append(
                        Document(
                            page_content=page_text,
                            metadata={
                                "filename": filename,
                                "page_number": page_index + 1,
                            },
                        )
                    )
        return documents

    @staticmethod
    def _parse_docx(filename: str, file_bytes: bytes) -> list[Document]:
        docx_document = DocxDocument(io.BytesIO(file_bytes))
        text_parts = [paragraph.text.strip() for paragraph in docx_document.paragraphs]

        # Tables frequently hold important clauses, dates, and payment terms.
        for table in docx_document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        document_text = "\n\n".join(part for part in text_parts if part)
        if not document_text:
            return []
        return [Document(page_content=document_text, metadata={"filename": filename})]

    @staticmethod
    def _parse_txt(filename: str, file_bytes: bytes) -> list[Document]:
        decoded_text = ""
        for encoding in ("utf-8-sig", "utf-16"):
            try:
                decoded_text = file_bytes.decode(encoding).strip()
                break
            except UnicodeDecodeError:
                continue

        if not decoded_text:
            raise DocumentProcessingError(
                "The TXT document must use UTF-8 or UTF-16 encoding."
            )
        return [Document(page_content=decoded_text, metadata={"filename": filename})]

    def _chat_with_document(
        self,
        message: str,
        document_id: str,
        history: list[BaseMessage],
        llm: ChatGoogleGenerativeAI,
    ) -> tuple[str, list[DocumentSource]]:
        retrieval_query = message
        if history:
            rewrite_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", QUESTION_REWRITE_SYSTEM_PROMPT),
                    MessagesPlaceholder("chat_history"),
                    ("human", "{input}"),
                ]
            )
            rewritten_question = (rewrite_prompt | llm | StrOutputParser()).invoke(
                {"input": message, "chat_history": history}
            )
            retrieval_query = rewritten_question.strip() or message

        retrieved_documents = self._retrieve_documents(
            document_id, retrieval_query, config.RETRIEVAL_COUNT
        )
        context = self._format_context(retrieved_documents)

        answer_prompt = ChatPromptTemplate.from_messages(
            [
                ("system", DOCUMENT_ANSWER_SYSTEM_PROMPT),
                MessagesPlaceholder("chat_history"),
                ("human", "{input}"),
            ]
        )
        answer = (answer_prompt | llm | StrOutputParser()).invoke(
            {"input": message, "chat_history": history, "context": context}
        )
        return self._prepare_answer_for_display(answer), self._collect_sources(
            retrieved_documents
        )

    @staticmethod
    def _general_chat(
        message: str,
        history: list[BaseMessage],
        llm: ChatGoogleGenerativeAI,
    ) -> str:
        general_prompt = ChatPromptTemplate.from_messages(
            [
                ("system", GENERAL_CHAT_SYSTEM_PROMPT),
                MessagesPlaceholder("chat_history"),
                ("human", "{input}"),
            ]
        )
        answer = (general_prompt | llm | StrOutputParser()).invoke(
            {"input": message, "chat_history": history}
        )
        return RAGService._prepare_answer_for_display(answer)

    def _retrieve_documents(
        self, document_id: str, retrieval_query: str, retrieval_count: int
    ) -> list[Document]:
        """Retrieve relevant chunks from one uploaded document's FAISS store."""
        document_store = self._get_document_store(document_id)
        retriever = document_store.vector_store.as_retriever(
            search_kwargs={"k": min(retrieval_count, document_store.chunk_count)}
        )
        return retriever.invoke(retrieval_query)

    def _get_document_store(self, document_id: str) -> DocumentStore:
        """Return one in-memory document store or raise a consistent error."""
        with self._stores_lock:
            document_store = self._document_stores.get(document_id)
        if document_store is None:
            raise UnknownDocumentError("The attached document is no longer available.")
        return document_store

    @staticmethod
    def _prepare_answer_for_display(answer: str) -> str:
        """Remove common Markdown markers that the plain-text chat UI cannot render."""
        plain_text = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", answer)
        plain_text = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", plain_text)
        plain_text = re.sub(r"(?m)^\s*\*\s+", "- ", plain_text)
        plain_text = plain_text.replace("**", "").replace("__", "").replace("`", "")
        plain_text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", plain_text)
        plain_text = re.sub(r"\n{3,}", "\n\n", plain_text)
        return plain_text.strip()

    @staticmethod
    def _to_langchain_history(history: list[ChatHistoryMessage]) -> list[BaseMessage]:
        recent_history = history[-config.MAX_CHAT_HISTORY_MESSAGES :]
        converted_history: list[BaseMessage] = []
        for history_message in recent_history:
            if history_message.role == "user":
                converted_history.append(HumanMessage(content=history_message.content))
            else:
                converted_history.append(AIMessage(content=history_message.content))
        return converted_history

    @staticmethod
    def _format_context(documents: list[Document]) -> str:
        if not documents:
            return "No relevant document context was retrieved."

        formatted_documents: list[str] = []
        for index, document in enumerate(documents, start=1):
            filename = document.metadata.get("filename", "Uploaded document")
            page_number = document.metadata.get("page_number")
            source_label = f"Source {index}: {filename}"
            if isinstance(page_number, int):
                source_label += f", page {page_number}"
            formatted_documents.append(
                f"{source_label}\nBEGIN DOCUMENT EXCERPT\n"
                f"{document.page_content}\nEND DOCUMENT EXCERPT"
            )
        return "\n\n".join(formatted_documents)

    @staticmethod
    def _collect_sources(documents: list[Document]) -> list[DocumentSource]:
        sources: list[DocumentSource] = []
        seen_sources: set[tuple[str, int | None]] = set()

        for document in documents:
            filename = document.metadata.get("filename")
            if not isinstance(filename, str) or not filename:
                continue
            page_number = document.metadata.get("page_number")
            if not isinstance(page_number, int):
                page_number = None

            source_key = (filename, page_number)
            if source_key in seen_sources:
                continue
            seen_sources.add(source_key)
            sources.append(DocumentSource(filename=filename, page_number=page_number))

        return sources
